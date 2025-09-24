#!/usr/bin/env python3
"""
Audio to Text Conversion Service
Supports multiple audio formats and languages including Spanish and English
"""

from flask import Flask, request, jsonify, render_template, send_from_directory, send_file
import speech_recognition as sr
import pydub
from pydub import AudioSegment
import os
import tempfile
import uuid
from werkzeug.utils import secure_filename
import logging
from typing import Optional, Dict, Any
import json
from docx import Document
from docx.shared import Inches
from datetime import datetime
import vosk
import wave
import json as json_lib
import numpy as np
# from pyannote.audio import Pipeline

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Production configuration
app.config['MAX_CONTENT_LENGTH'] = int(os.environ.get('MAX_CONTENT_LENGTH', 25 * 1024 * 1024))  # 25MB max file size
app.config['UPLOAD_FOLDER'] = os.environ.get('UPLOAD_FOLDER', tempfile.gettempdir())
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')

# Supported audio formats
SUPPORTED_FORMATS = {
    'mp3': 'mp3',
    'wav': 'wav', 
    'aac': 'aac',
    'm4a': 'm4a',
    'ogg': 'ogg',
    'flac': 'flac',
    'wma': 'wma',
    'mp4': 'mp4',
    'webm': 'webm'
}

# Language mapping for speech recognition
# Language mapping for Vosk models (only Spanish and English to save memory)
LANGUAGE_MAPPING = {
    'spanish': 'es-ES',
    'english': 'en-US',
    'auto': None
}

# Vosk model paths (only Spanish and English)
VOSK_MODELS = {
    'es-ES': '/app/models/vosk-model-small-es-0.42',
    'en-US': '/app/models/vosk-model-small-en-us-0.15'
}

class AudioToTextConverter:
    def __init__(self):
        self.recognizer = sr.Recognizer()
        self.recognizer.energy_threshold = 300
        self.recognizer.dynamic_energy_threshold = True
        self.recognizer.pause_threshold = 0.8
        self.recognizer.operation_timeout = None
        self.recognizer.phrase_threshold = 0.3
        self.recognizer.non_speaking_duration = 0.8

        # Cache for loaded Vosk models
        self.models = {}
        
        # Speaker diarization disabled to save memory
        # self.diarization_pipeline = None
        # self._load_diarization_pipeline()
    
    def _get_vosk_model(self, lang_code: str):
        """Load or return cached Vosk model for given language code"""
        if lang_code not in VOSK_MODELS:
            raise Exception(f"No Vosk model available for language {lang_code}")

        if lang_code not in self.models:
            model_path = VOSK_MODELS[lang_code]
            if not os.path.exists(model_path):
                raise Exception(f"Model path not found: {model_path}")
            logger.info(f"Loading Vosk model for {lang_code} from {model_path}")
            self.models[lang_code] = vosk.Model(model_path)
            logger.info(f"Vosk model for {lang_code} loaded successfully")

        return self.models[lang_code]
    
    # def _load_diarization_pipeline(self):
    #     """Load speaker diarization pipeline"""
    #     try:
    #         logger.info("Loading speaker diarization pipeline...")
    #         # Use a different model that doesn't require authentication
    #         self.diarization_pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization")
    #         logger.info("Speaker diarization pipeline loaded successfully")
    #     except Exception as e:
    #         logger.warning(f"Failed to load speaker diarization pipeline: {e}")
    #         logger.warning("Speaker diarization will not be available")
    #         self.diarization_pipeline = None
    
    # def _perform_speaker_diarization(self, audio_path: str) -> list:
    #     """Perform speaker diarization on audio file"""
    #     if not self.diarization_pipeline:
    #         logger.warning("Speaker diarization pipeline not available")
    #         return []
    #     
    #     try:
    #         logger.info("Performing speaker diarization...")
    #         diarization = self.diarization_pipeline(audio_path)
    #         
    #         speaker_segments = []
    #         for turn, _, speaker in diarization.itertracks(yield_label=True):
    #             speaker_segments.append({
    #                 'speaker': speaker,
    #                 'start': turn.start,
    #                 'end': turn.end,
    #                 'duration': turn.end - turn.start
    #             })
    #         
    #         logger.info(f"Found {len(speaker_segments)} speaker segments")
    #         return speaker_segments
    #         
    #     except Exception as e:
    #         logger.error(f"Speaker diarization failed: {e}")
    #         return []
        
    def convert_audio_to_wav(self, input_path: str, output_path: str) -> bool:
        """Convert any supported audio format to WAV for processing"""
        try:
            # Get file extension
            file_ext = os.path.splitext(input_path)[1].lower().lstrip('.')
            
            if file_ext == 'wav':
                # Already WAV, just copy
                import shutil
                shutil.copy2(input_path, output_path)
                return True
            
            # Load audio with pydub
            if file_ext in ['mp3', 'aac', 'm4a', 'ogg', 'flac', 'wma', 'mp4', 'webm']:
                audio = AudioSegment.from_file(input_path, format=file_ext)
            else:
                # Try to load with pydub's auto-detection
                audio = AudioSegment.from_file(input_path)
            
            # Convert to WAV format
            audio.export(output_path, format="wav")
            logger.info(f"Successfully converted {input_path} to WAV")
            return True
            
        except Exception as e:
            logger.error(f"Error converting audio to WAV: {str(e)}")
            return False
    
    def transcribe_audio(self, audio_path: str, language: str = 'auto') -> Dict[str, Any]:
        """Transcribe audio file to text"""
        result = {
            'success': False,
            'text': '',
            'confidence': 0.0,
            'language_detected': None,
            'error': None
        }
        
        try:
            # Convert to WAV if needed
            wav_path = audio_path
            if not audio_path.lower().endswith('.wav'):
                wav_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}.wav")
                if not self.convert_audio_to_wav(audio_path, wav_path):
                    result['error'] = "Failed to convert audio to WAV format"
                    return result
            
            # Load audio file
            with sr.AudioFile(wav_path) as source:
                # Adjust for ambient noise
                self.recognizer.adjust_for_ambient_noise(source, duration=0.5)
                audio_data = self.recognizer.record(source)
            
            # Get language code
            lang_code = LANGUAGE_MAPPING.get(language.lower(), None)
            
            # Use only offline recognition (Vosk)
            # Check if language is supported
            if language.lower() not in LANGUAGE_MAPPING:
                result['error'] = f"Language '{language}' is not supported. Supported languages: {', '.join(LANGUAGE_MAPPING.keys())}"
                result['supported_languages'] = list(LANGUAGE_MAPPING.keys())
                result['note'] = 'For other languages, consider using online speech recognition services.'
            else:
                try:
                    # Get the appropriate model for the language
                    if lang_code is None:  # auto-detect, default to English
                        lang_code = 'en-US'
                    
                    model = self._get_vosk_model(lang_code)
                    text, confidence, word_timestamps = self._transcribe_with_vosk(audio_data, model)
                    
                    if text and text.strip():
                        result['success'] = True
                        result['text'] = text.strip()
                        result['confidence'] = confidence
                        result['service_used'] = 'vosk'
                        result['language_detected'] = lang_code
                        result['note'] = f'Transcribed using Vosk model for {lang_code}'
                        result['word_timestamps'] = word_timestamps
                    else:
                        result['error'] = "No speech detected in audio. Try speaking more clearly or check if the audio contains speech."
                except Exception as e:
                    logger.error(f"Vosk recognition failed for {language}: {str(e)}")
                    result['error'] = f"Speech recognition failed: {str(e)}"
            
            # Only set generic error if no specific error was set
            if not result['success'] and 'error' not in result:
                result['error'] = "All speech recognition services failed"
            
            # Clean up temporary WAV file
            if wav_path != audio_path and os.path.exists(wav_path):
                os.remove(wav_path)
                
        except Exception as e:
            logger.error(f"Error in transcription: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _transcribe_with_vosk(self, audio_data, model) -> tuple:
        """Transcribe using Vosk (offline only) with timestamps - optimized for large files"""
        try:
            # Log audio data info for debugging
            logger.info(f"Audio data: {len(audio_data.frame_data)} bytes, {audio_data.sample_rate}Hz, {audio_data.sample_width} bytes per sample")

            # Create Vosk recognizer
            rec = vosk.KaldiRecognizer(model, audio_data.sample_rate)
            rec.SetWords(True)

            # Process audio data
            logger.info("Attempting Vosk recognition with timestamps")

            # Convert audio data to bytes
            audio_bytes = audio_data.frame_data

            # Use larger chunks for better performance with large files
            chunk_size = 8000  # Increased from 4000
            results = []
            word_timestamps = []
            processed_chunks = 0
            total_chunks = len(audio_bytes) // chunk_size + 1

            for i in range(0, len(audio_bytes), chunk_size):
                chunk = audio_bytes[i:i + chunk_size]
                if rec.AcceptWaveform(chunk):
                    result = json_lib.loads(rec.Result())
                    if 'text' in result and result['text']:
                        results.append(result['text'])
                    # Extract word-level timestamps if available
                    if 'result' in result:
                        for word_info in result['result']:
                            word_timestamps.append({
                                'word': word_info.get('word', ''),
                                'start': word_info.get('start', 0),
                                'end': word_info.get('end', 0),
                                'conf': word_info.get('conf', 0)
                            })
                
                processed_chunks += 1
                # Log progress every 100 chunks to avoid spam
                if processed_chunks % 100 == 0:
                    logger.info(f"Processed {processed_chunks}/{total_chunks} chunks ({processed_chunks/total_chunks*100:.1f}%)")

            # Get final result
            final_result = json_lib.loads(rec.FinalResult())
            if 'text' in final_result and final_result['text']:
                results.append(final_result['text'])
            # Extract final word timestamps
            if 'result' in final_result:
                for word_info in final_result['result']:
                    word_timestamps.append({
                        'word': word_info.get('word', ''),
                        'start': word_info.get('start', 0),
                        'end': word_info.get('end', 0),
                        'conf': word_info.get('conf', 0)
                    })

            # Combine all results
            full_text = ' '.join(results).strip()

            if full_text:
                # Calculate overall confidence from word confidences
                if word_timestamps:
                    avg_confidence = sum(w['conf'] for w in word_timestamps) / len(word_timestamps)
                else:
                    avg_confidence = min(0.8, 0.5 + (len(full_text.split()) * 0.05))

                # Log the results for debugging
                logger.info(f"Transcription completed: {len(full_text)} characters, {len(word_timestamps)} word timestamps")

                return full_text, avg_confidence, word_timestamps
            else:
                raise Exception("No speech detected in audio")

        except Exception as e:
            raise Exception(f"Vosk recognition error: {e}")
    
    def _transcribe_with_speakers(self, audio_path: str, language: str) -> Dict[str, Any]:
        """Transcribe audio with speaker diarization (disabled - falls back to regular transcription)"""
        logger.info("Speaker diarization is disabled, falling back to regular transcription")
        return self.transcribe_audio(audio_path, language)
        
        try:
            # Get language code
            lang_code = LANGUAGE_MAPPING.get(language.lower(), 'en-US')
            if lang_code is None:
                lang_code = 'en-US'
            
            # Perform speaker diarization
            speaker_segments = self._perform_speaker_diarization(audio_path)
            
            if not speaker_segments:
                logger.warning("No speaker segments found, falling back to regular transcription")
                # Fall back to regular transcription
                return self.transcribe_audio(audio_path, language)
            
            # Load the appropriate Vosk model
            model = self._get_vosk_model(lang_code)
            
            # Transcribe each speaker segment
            transcribed_segments = []
            total_confidence = 0.0
            segment_count = 0
            
            for segment in speaker_segments:
                try:
                    # Extract audio segment
                    audio = AudioSegment.from_wav(audio_path)
                    start_ms = int(segment['start'] * 1000)
                    end_ms = int(segment['end'] * 1000)
                    segment_audio = audio[start_ms:end_ms]
                    
                    # Save segment to temporary file
                    segment_path = f"/tmp/segment_{uuid.uuid4().hex}.wav"
                    segment_audio.export(segment_path, format="wav")
                    
                    # Transcribe segment
                    with sr.AudioFile(segment_path) as source:
                        audio_data = self.recognizer.record(source)
                    
                    text, confidence, word_timestamps = self._transcribe_with_vosk(audio_data, model)
                    
                    if text and text.strip():
                        transcribed_segments.append({
                            'speaker': segment['speaker'],
                            'start_time': segment['start'],
                            'end_time': segment['end'],
                            'duration': segment['duration'],
                            'text': text.strip(),
                            'confidence': confidence,
                            'word_timestamps': word_timestamps
                        })
                        total_confidence += confidence
                        segment_count += 1
                    
                    # Clean up temporary file
                    os.remove(segment_path)
                    
                except Exception as e:
                    logger.warning(f"Failed to transcribe segment for {segment['speaker']}: {e}")
                    continue
            
            if transcribed_segments:
                result['success'] = True
                result['speaker_segments'] = transcribed_segments
                result['confidence'] = total_confidence / segment_count if segment_count > 0 else 0.0
                result['language_detected'] = lang_code
                result['text'] = '\n\n'.join([f"{seg['speaker']}: {seg['text']}" for seg in transcribed_segments])
                result['note'] = f'Transcribed with speaker diarization using Vosk model for {lang_code}'
            else:
                result['error'] = "No speech detected in any speaker segments"
                
        except Exception as e:
            logger.error(f"Speaker-based transcription failed: {e}")
            result['error'] = f"Speaker-based transcription failed: {str(e)}"
        
        return result

def create_word_document(transcription_result: Dict[str, Any], original_filename: str) -> str:
    """Create a Word document with the transcription results"""
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Audio Transcription Report', 0)
        
        # Add metadata section
        doc.add_heading('Document Information', level=1)
        
        # Create a table for metadata
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Property'
        hdr_cells[1].text = 'Value'
        
        # Add metadata rows
        metadata = [
            ('Original File', original_filename),
            ('Transcription Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Service Used', transcription_result.get('service_used', 'Unknown')),
            ('Language Detected', transcription_result.get('language_detected', 'Unknown')),
            ('Confidence Score', f"{transcription_result.get('confidence', 0):.2f}"),
            ('File Size', f"{transcription_result.get('file_size', 0) / 1024 / 1024:.2f} MB")
        ]
        
        for prop, value in metadata:
            row_cells = table.add_row().cells
            row_cells[0].text = prop
            row_cells[1].text = str(value)
        
        # Add transcription section
        doc.add_heading('Transcription', level=1)
        
        if transcription_result.get('success'):
            # Check if this is speaker-based transcription
            speaker_segments = transcription_result.get('speaker_segments', [])
            
            if speaker_segments:
                # Add speaker-based transcription
                doc.add_heading('Speaker-Based Transcription', level=2)
                
                for segment in speaker_segments:
                    # Add speaker header
                    speaker_para = doc.add_paragraph()
                    speaker_para.add_run(f"{segment['speaker']} ").bold = True
                    speaker_para.add_run(f"({segment['start_time']:.1f}s - {segment['end_time']:.1f}s, {segment['duration']:.1f}s)")
                    
                    # Add transcribed text
                    doc.add_paragraph(segment['text'])
                    
                    # Add confidence for this segment
                    doc.add_paragraph(f"Confidence: {segment['confidence']:.2f}")
                    doc.add_paragraph()  # Add spacing
                
                # Add summary table
                doc.add_heading('Speaker Summary', level=2)
                summary_table = doc.add_table(rows=1, cols=4)
                summary_table.style = 'Table Grid'
                hdr_cells = summary_table.rows[0].cells
                hdr_cells[0].text = 'Speaker'
                hdr_cells[1].text = 'Duration (s)'
                hdr_cells[2].text = 'Words'
                hdr_cells[3].text = 'Avg Confidence'
                
                for segment in speaker_segments:
                    row_cells = summary_table.add_row().cells
                    row_cells[0].text = segment['speaker']
                    row_cells[1].text = f"{segment['duration']:.1f}"
                    row_cells[2].text = str(len(segment['text'].split()))
                    row_cells[3].text = f"{segment['confidence']:.2f}"
            else:
                # Add regular transcription
                doc.add_paragraph(transcription_result.get('text', ''))
            
            # Add note if available
            if transcription_result.get('note'):
                doc.add_paragraph(f"Note: {transcription_result.get('note')}")
        else:
            # Add error message
            error_para = doc.add_paragraph()
            error_para.add_run('Transcription Failed: ').bold = True
            error_para.add_run(transcription_result.get('error', 'Unknown error'))
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph('Generated by Audio to Text Conversion Service (Offline)')
        footer.alignment = 1  # Center alignment
        
        # Save to temporary file
        temp_filename = f"transcription_{uuid.uuid4().hex}.docx"
        temp_path = os.path.join(tempfile.gettempdir(), temp_filename)
        doc.save(temp_path)
        
        logger.info(f"Word document created: {temp_path}")
        return temp_path
        
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        raise Exception(f"Failed to create Word document: {str(e)}")

# Initialize converter
converter = AudioToTextConverter()

@app.route('/')
def index():
    """Main page with upload form"""
    return render_template('index.html', 
                         supported_formats=list(SUPPORTED_FORMATS.keys()),
                         languages=list(LANGUAGE_MAPPING.keys()))

@app.route('/api/convert', methods=['POST'])
def convert_audio():
    """API endpoint for audio to text conversion"""
    try:
        # Check if file is present
        if 'audio_file' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        file = request.files['audio_file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check file size (limit to 25MB to prevent memory issues)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > 25 * 1024 * 1024:  # 25MB limit
            return jsonify({'error': 'File too large. Maximum size is 25MB. Please use a smaller audio file.'}), 400
        
        # Get language parameter
        language = request.form.get('language', 'auto')
        
        # Validate file format
        filename = secure_filename(file.filename)
        file_ext = os.path.splitext(filename)[1].lower().lstrip('.')
        
        if file_ext not in SUPPORTED_FORMATS:
            return jsonify({
                'error': f'Unsupported file format: {file_ext}. Supported formats: {", ".join(SUPPORTED_FORMATS.keys())}'
            }), 400
        
        # Save uploaded file
        temp_filename = f"{uuid.uuid4()}.{file_ext}"
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename)
        file.save(temp_path)
        
        try:
            # Convert audio to text
            result = converter.transcribe_audio(temp_path, language)
            
            # Add metadata
            result['filename'] = filename
            result['file_size'] = os.path.getsize(temp_path)
            result['language_requested'] = language
            
            # Create Word document
            doc_path = create_word_document(result, filename)
            
            # Generate download filename
            base_name = os.path.splitext(filename)[0]
            download_filename = f"{base_name}_transcription.docx"
            
            # Return the Word document as download
            return send_file(
                doc_path,
                as_attachment=True,
                download_name=download_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        finally:
            # Clean up uploaded file
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    except Exception as e:
        logger.error(f"Error in convert_audio: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/formats')
def get_supported_formats():
    """Get list of supported audio formats"""
    return jsonify({
        'formats': list(SUPPORTED_FORMATS.keys()),
        'languages': list(LANGUAGE_MAPPING.keys())
    })

@app.route('/api/health')
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'service': 'audio-to-text-converter',
        'supported_formats': len(SUPPORTED_FORMATS),
        'supported_languages': len(LANGUAGE_MAPPING)
    })


@app.route('/api/test-offline')
def test_offline_recognition():
    """Test offline speech recognition capabilities"""
    try:
        # Test if we can load the English model
        try:
            model = converter._get_vosk_model('en-US')
            model_loaded = True
        except Exception as e:
            return jsonify({
                'status': 'error', 
                'message': f'Failed to load English model: {str(e)}',
                'available_languages': list(LANGUAGE_MAPPING.keys())
            })
        
        # Create a simple test audio (1 second of silence)
        import io
        import wave
        
        # Create a simple WAV file with silence
        buffer = io.BytesIO()
        with wave.open(buffer, 'wb') as wav_file:
            wav_file.setnchannels(1)  # Mono
            wav_file.setsampwidth(2)  # 16-bit
            wav_file.setframerate(16000)  # 16kHz
            wav_file.writeframes(b'\x00' * 16000)  # 1 second of silence
        
        buffer.seek(0)
        
        # Try to recognize the silence (should return empty result)
        with sr.AudioFile(buffer) as source:
            audio_data = converter.recognizer.record(source)
        
        try:
            # Test Vosk recognition with silence
            text, confidence = converter._transcribe_with_vosk(audio_data, model)
            if not text or text.strip() == "":
                return jsonify({
                    'status': 'success', 
                    'message': 'Vosk is working correctly (no speech detected in silence)',
                    'available_languages': list(LANGUAGE_MAPPING.keys())
                })
            else:
                return jsonify({
                    'status': 'warning', 
                    'message': f'Vosk detected speech in silence: "{text}"',
                    'available_languages': list(LANGUAGE_MAPPING.keys())
                })
        except Exception as e:
            if "No speech detected" in str(e):
                return jsonify({
                    'status': 'success', 
                    'message': 'Vosk is working correctly (no speech detected in silence)',
                    'available_languages': list(LANGUAGE_MAPPING.keys())
                })
            else:
                return jsonify({
                    'status': 'error', 
                    'message': f'Vosk error: {e}',
                    'available_languages': list(LANGUAGE_MAPPING.keys())
                })
            
    except Exception as e:
        return jsonify({
            'status': 'error', 
            'message': f'Test failed: {e}',
            'available_languages': list(LANGUAGE_MAPPING.keys())
        })

if __name__ == '__main__':
    # Create templates directory if it doesn't exist
    os.makedirs('templates', exist_ok=True)
    os.makedirs('static', exist_ok=True)
    
    # Get port from environment (for Render.com deployment)
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV', 'production') == 'development'
    
    print("🎵 Audio to Text Conversion Service (OFFLINE ONLY)")
    print("=" * 50)
    print(f"Supported formats: {', '.join(SUPPORTED_FORMATS.keys())}")
    print(f"Supported languages: {', '.join(LANGUAGE_MAPPING.keys())}")
    print("⚠️  NOTE: Vosk only supports English")
    print("=" * 50)
    print(f"Starting server on port {port}")
    print(f"Debug mode: {debug}")
    
    app.run(debug=debug, host='0.0.0.0', port=port)
